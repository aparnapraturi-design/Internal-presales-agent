from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Literal, Optional

from pypdf import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import markdown as mdlib



MDMode = Literal["keep", "plain"]


@dataclass
class ExtractOptions:
    """
    Controls extraction behavior and safety limits.
    """
    max_bytes: int = 25 * 1024 * 1024  # 25 MB
    md_mode: MDMode = "keep"
    html_remove_scripts: bool = True
    html_collapse_whitespace: bool = True


class ExtractionError(RuntimeError):
    pass


def _ensure_readable_file(path: Path, opts: ExtractOptions) -> None:
    if not path.exists() or not path.is_file():
        raise ExtractionError(f"File not found: {path}")
    size = path.stat().st_size
    if size > opts.max_bytes:
        raise ExtractionError(f"File too large ({size} bytes). Max allowed: {opts.max_bytes} bytes. File: {path}")


# ============================================================
# PDF
# ============================================================

def extract_text_pdf(path: Path, opts: ExtractOptions) -> str:
    _ensure_readable_file(path, opts)
    try:
        reader = PdfReader(str(path))
        parts: list[str] = []
        for i, page in enumerate(reader.pages):
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            txt = txt.strip()
            if txt:
                parts.append(txt)
        return "\n\n".join(parts).strip()
    except Exception as e:
        raise ExtractionError(f"PDF extraction failed for {path.name}: {type(e).__name__}: {e}") from e


# ============================================================
# DOCX
# ============================================================

def extract_text_docx(path: Path, opts: ExtractOptions) -> str:
    _ensure_readable_file(path, opts)
    try:
        doc = DocxDocument(str(path))
        parts: list[str] = []

        # paragraphs
        for para in doc.paragraphs:
            t = (para.text or "").strip()
            if t:
                parts.append(t)

        # tables (rows with | separators)
        for t in doc.tables:
            for row in t.rows:
                cells = [ (c.text or "").strip().replace("\n", " ") for c in row.cells ]
                if any(cells):
                    # de-dupe repeated cells sometimes seen in merged cells
                    line = " | ".join(cells)
                    line = re.sub(r"\s+\|\s+", " | ", line).strip()
                    parts.append(line)

        return "\n".join(parts).strip()
    except Exception as e:
        raise ExtractionError(f"DOCX extraction failed for {path.name}: {type(e).__name__}: {e}") from e


# ============================================================
# HTML
# ============================================================

def extract_text_html(path: Path, opts: ExtractOptions) -> str:
    _ensure_readable_file(path, opts)

    # Read as bytes -> decode leniently
    raw = path.read_bytes()
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        # fallback latin-1-ish
        text = raw.decode("utf-8", errors="ignore")

    try:
        soup = BeautifulSoup(text, "lxml")

        if opts.html_remove_scripts:
            for tag in soup(["script", "style", "noscript"]):
                tag.decompose()

        out = soup.get_text(separator="\n")

        # clean up whitespace
        out = out.replace("\r\n", "\n").replace("\r", "\n")
        if opts.html_collapse_whitespace:
            out = re.sub(r"[ \t]+", " ", out)
            out = re.sub(r"\n{3,}", "\n\n", out)

        return out.strip()
    except Exception as e:
        raise ExtractionError(f"HTML extraction failed for {path.name}: {type(e).__name__}: {e}") from e


# ============================================================
# Markdown
# ============================================================

def extract_markdown_keep(path: Path, opts: ExtractOptions) -> str:
    _ensure_readable_file(path, opts)
    return path.read_text(encoding="utf-8", errors="ignore").strip()

def extract_markdown_plain(path: Path, opts: ExtractOptions) -> str:
    _ensure_readable_file(path, opts)
    md_text = path.read_text(encoding="utf-8", errors="ignore")
    html = mdlib.markdown(md_text, extensions=["tables", "fenced_code"])
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    out = soup.get_text(separator="\n")
    out = out.replace("\r\n", "\n").replace("\r", "\n")
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()


# ============================================================
# TXT
# ============================================================

def extract_text_txt(path: Path, opts: ExtractOptions) -> str:
    _ensure_readable_file(path, opts)
    return path.read_text(encoding="utf-8", errors="ignore").strip()

#--------------------------------------------------------------
# Audipio transcription with Whisper
#--------------------------------------------------------------

def transcribe_with_whisper(file_path: Path, *, model_size: str, language: Optional[str]) -> str:
    """
    Optional audio path. Only imports whisper when needed.
    """
    try:
        from faster_whisper import WhisperModel
    except Exception as e:
        raise RuntimeError(
            "Audio transcription requested but faster-whisper/ctranslate2 is not working on this machine. "
            f"Import error: {e}"
        )

    model = WhisperModel(model_size, device="cpu", compute_type="int8")
    segments, _info = model.transcribe(str(file_path), language=language)
    return " ".join(seg.text.strip() for seg in segments if seg.text).strip()


# ============================================================
# Dispatcher
# ============================================================

def extract_text_any(path: str | Path, opts: Optional[ExtractOptions] = None) -> str:
    opts = opts or ExtractOptions()
    p = Path(path)
    ext = p.suffix.lower()

    if ext == ".pdf":
        return extract_text_pdf(p, opts)
    if ext == ".docx":
        return extract_text_docx(p, opts)
    if ext in (".html", ".htm"):
        return extract_text_html(p, opts)
    if ext == ".md":
        if opts.md_mode == "plain":
            return extract_markdown_plain(p, opts)
        return extract_markdown_keep(p, opts)
    if ext == ".txt":
        return extract_text_txt(p, opts)
    if ext in (".mp3", ".wav", ".m4a", ".flac", ".aac", ".ogg"):
        return transcribe_with_whisper(file_path=p, model_size="small", language="en")

    raise ExtractionError(f"Unsupported file type for extraction: {ext} ({p.name})")
